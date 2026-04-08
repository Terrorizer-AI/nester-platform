"""
SOW Generator API.

Three document types:
  - reference_sow : Previous SOW → AI copies the SECTION STRUCTURE (headings, order)
  - template      : Design template DOCX → used for VISUAL DESIGN (fonts, colors, logos)
  - proposal      : Client proposals → AI fills in NEW CONTENT from these

Endpoints:
  POST   /sow/sessions                            → Create session
  GET    /sow/sessions                            → List sessions
  GET    /sow/sessions/{id}                       → Get session detail + chat
  DELETE /sow/sessions/{id}                       → Delete session + cascade
  PATCH  /sow/sessions/{id}                       → Update title
  POST   /sow/reference-sows                      → Upload reference SOW (global)
  GET    /sow/reference-sows                      → List reference SOWs
  DELETE /sow/reference-sows/{doc_id}             → Remove reference SOW
  POST   /sow/templates                           → Upload design template (global)
  GET    /sow/templates                           → List design templates
  DELETE /sow/templates/{doc_id}                  → Remove design template
  POST   /sow/sessions/{id}/proposals             → Upload proposals
  GET    /sow/sessions/{id}/proposals             → List proposals
  DELETE /sow/sessions/{id}/proposals/{doc_id}    → Remove proposal
  POST   /sow/sessions/{id}/generate              → Initial SOW generation (SSE)
  POST   /sow/sessions/{id}/chat                  → Chat refinement (SSE)
  PUT    /sow/sessions/{id}/sow                   → Direct markdown edit
  GET    /sow/sessions/{id}/preview               → Live HTML preview of the DOCX
  GET    /sow/sessions/{id}/download              → Download as DOCX
"""

from __future__ import annotations

import io
import json
import logging
import re
import uuid
from typing import Any

from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import HTMLResponse, StreamingResponse
from pydantic import BaseModel

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/sow", tags=["sow"])

ALLOWED_PROPOSAL_EXT = {".docx", ".pdf", ".pptx"}


# ── Helpers ──────────────────────────────────────────────────────────────────


def _extract_text(raw: bytes, filename: str) -> str:
    """Extract text from a document based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        from knowledge.drive_sync import _extract_pdf_text
        return _extract_pdf_text(raw)
    if ext == "docx":
        from knowledge.drive_sync import _extract_docx_text
        return _extract_docx_text(raw)
    if ext == "pptx":
        from knowledge.drive_sync import _extract_pptx_text
        return _extract_pptx_text(raw)
    if ext in ("txt", "md", "csv"):
        return raw.decode("utf-8", errors="replace")
    return ""


def _new_id() -> str:
    return uuid.uuid4().hex[:16]


def _doc_id_for(filename: str) -> str:
    """Deterministic doc ID so re-uploading same name replaces it."""
    return str(uuid.uuid5(uuid.NAMESPACE_URL, filename))


# ── Pydantic Models ─────────────────────────────────────────────────────────


class CreateSessionReq(BaseModel):
    title: str = "Untitled SOW"


class UpdateTitleReq(BaseModel):
    title: str


class SOWChatReq(BaseModel):
    message: str


class UpdateSOWReq(BaseModel):
    markdown: str


# ── Session CRUD ─────────────────────────────────────────────────────────────


@router.post("/sessions")
async def create_session(req: CreateSessionReq):
    from memory.sqlite_ops import create_sow_session, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    sid = _new_id()
    create_sow_session(sid, req.title)
    return {"id": sid, "title": req.title}


@router.get("/sessions")
async def list_sessions():
    from memory.sqlite_ops import list_sow_sessions, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    return {"sessions": list_sow_sessions()}


@router.get("/sessions/{session_id}")
async def get_session(session_id: str):
    from memory.sqlite_ops import get_sow_session, list_sow_chat_messages, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    session = get_sow_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    messages = list_sow_chat_messages(session_id)
    return {
        "session": session,
        "messages": messages,
    }


@router.delete("/sessions/{session_id}")
async def delete_session(session_id: str):
    from memory.sqlite_ops import delete_sow_session, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    delete_sow_session(session_id)
    return {"ok": True}


@router.patch("/sessions/{session_id}")
async def update_session_title(session_id: str, req: UpdateTitleReq):
    from memory.sqlite_ops import update_sow_session_title, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    update_sow_session_title(session_id, req.title)
    return {"ok": True}


# ── Proposal CRUD ────────────────────────────────────────────────────────────


@router.post("/sessions/{session_id}/proposals")
async def upload_proposals(session_id: str, files: list[UploadFile] = File(...)):
    from memory.sqlite_ops import save_sow_document, get_sow_session, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    if not get_sow_session(session_id):
        raise HTTPException(404, "Session not found")
    saved: list[dict[str, Any]] = []
    for f in files:
        ext = f.filename.rsplit(".", 1)[-1].lower() if f.filename and "." in f.filename else ""
        if f".{ext}" not in ALLOWED_PROPOSAL_EXT:
            raise HTTPException(400, f"Proposals must be .docx, .pdf, or .pptx, got: {f.filename}")
        raw = await f.read()
        text = _extract_text(raw, f.filename or "file.docx")
        doc_id = _doc_id_for(f"{session_id}_{f.filename or 'file'}")
        save_sow_document(
            doc_id=doc_id, session_id=session_id, doc_type="proposal",
            file_name=f.filename or "file",
            mime_type=f.content_type or "application/octet-stream",
            raw_bytes=raw, extracted_text=text,
        )
        saved.append({"id": doc_id, "file_name": f.filename, "chars": len(text)})
    return {"uploaded": saved}


@router.get("/sessions/{session_id}/proposals")
async def list_proposals(session_id: str):
    from memory.sqlite_ops import list_sow_documents, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    return {"proposals": list_sow_documents(session_id=session_id, doc_type="proposal")}


@router.delete("/sessions/{session_id}/proposals/{doc_id}")
async def delete_proposal(session_id: str, doc_id: str):
    from memory.sqlite_ops import delete_sow_document, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    delete_sow_document(doc_id)
    return {"ok": True}


# ── SOW Generation (SSE) ────────────────────────────────────────────────────


GENERATE_SYSTEM = """You are a senior proposal writer and legal document specialist at Nester Labs, with 15+ years of experience writing Statements of Work for Fortune 500 consulting engagements. You write with the precision of a lawyer, the clarity of a business analyst, and the persuasiveness of a sales executive.

You generate SOWs for **Nester Labs India Private Limited** (signing authority: Kunal Srivastava, Director).

You will receive **CLIENT PROPOSALS** — the ONLY source of all project details.

CRITICAL DATA RULES:
- ALL content MUST come directly from the uploaded proposals — client name, project name, scope, deliverables, timelines, team, budget, requirements, milestones, contact info, payment terms, assumptions
- Extract every relevant detail — names, dates, numbers, deliverables, milestones, constraints, technologies, roles
- Do NOT invent, assume, or fabricate ANY data that is not in the proposals
- Do NOT use generic placeholder content like "Company X" or "the client" — use the actual client name from the proposals
- Do NOT use example dates, budgets, or deliverables — use only what the proposals state
- If a specific detail is NOT in the proposals, mark it as [REVIEW NEEDED: what is missing] — NEVER fill it with made-up data
- Expand proposal content into full professional prose, but the underlying facts must all trace back to the uploaded documents

You MUST use the following FIXED section structure for every SOW. Do not deviate from this order or these headings:

```
# 1. Executive Summary
   (2-3 paragraphs: what the project is, who the parties are, engagement duration, governed by which MSA)

# 2. Scope of Work
   (detailed breakdown of everything the Vendor will build/deliver)
   ## 2.1 [subsections per functional area — e.g. Infrastructure, AI Services, etc.]

# 3. Project Timeline, Milestones, and Acceptance Criteria
   ## 3.1 Project Timeline
       (Start date, end date, total weeks)
   ## 3.2 Milestones and Acceptance Criteria
       (table: Milestone | Deliverable | Due Date | Owner + acceptance criteria per milestone)

# 4. Billing and Payment Terms
   ## 4.1 Model and Duration
       (Fixed Bid / T&M, duration, total cost)
   ## 4.2 Payment Schedule
       (table: Phase | Amount | Due Date | Trigger)
   ## 4.3 Payment Terms
       (net-7/net-30, invoice routing, billing contact)

# 5. Assumptions
   (bullet list of all project assumptions)

# 6. Licenses, Tools, and Infrastructure
   ## 6.1 Licenses Provided by Nester Labs
   ## 6.2 Client-Managed Infrastructure

# 7. Points of Contact
   ## 7.1 Nester Labs Point of Contact
   ## 7.2 [Client Name] Point of Contact

# 8. Signatures of Both Parties
   (DO NOT write signature lines — they are auto-generated in the DOCX. Just write:
    "This SOW is executed by the duly authorized representatives of both parties as of the last signature date below.")
```

IMPORTANT: Section 8 (Signatures) must ONLY contain the single sentence above. The actual signature block is programmatically inserted into the DOCX — do NOT create signature lines, underscores, name fields, or date fields in markdown.

---

WRITING STANDARDS — follow these rigorously:

**Length & Completeness**
- Every section must be FULLY written — no placeholders unless data is genuinely missing
- Match or exceed the length of the reference SOW section-by-section
- Each body section should have at minimum 2-4 substantive paragraphs
- Use specific numbers, dates, names, and metrics from the proposals — never be vague

**Humanized, Non-AI Writing Style**
- Write in confident, active voice: "The Vendor will deliver..." not "It is expected that delivery will occur..."
- Vary sentence length — mix short punchy sentences with longer detailed ones
- Use industry-specific terminology naturally, not robotically
- Avoid filler phrases like "It is important to note", "In order to", "As mentioned above"
- Write as if a senior consultant personally drafted this for this specific client

**Swiss / International Typographic Style — Grid Layout**

You MUST follow this layout pattern for EVERY section. The document uses a strict
Swiss grid system — every heading is paired with body content in a predictable,
rhythmic pattern. Consistency is paramount: every section must look identical in
structure and spacing.

Grid rules:
- EVERY section follows the exact same pattern:
  1. Section heading (H1/H2/H3)
  2. One blank line
  3. Body paragraphs, tables, or bullet lists
  4. One blank line before the next heading
- NO orphan headings — every heading MUST have at least one paragraph beneath it
- NO double headings — never put two headings back-to-back without body text between them
- Body text under a heading must be at minimum 2 sentences (never a single dangling sentence)

Section numbering pattern:
- H1 sections: numbered sequentially — "1. Executive Summary", "2. Scope of Work", "3. Deliverables"
- H2 subsections: numbered as parent.child — "2.1 Technical Requirements", "2.2 Integration Scope"
- H3 sub-subsections: parent.child.child — "2.1.1 API Endpoints"
- The section number is part of the heading text itself (not generated separately)

Paragraph rhythm (Swiss grid):
- Each section opens with a lead paragraph that summarizes the section in 2-3 sentences
- Followed by detailed paragraphs expanding each point
- Close with a transition sentence or summary statement connecting to the next section
- Maintain uniform paragraph density across sections — if Section 2 has 4 paragraphs, Section 3 should have 3-5 (never just 1)

**Tables — Use Extensively**
- Milestones → table with columns: Milestone | Deliverable | Due Date | Owner
- Deliverables → table with columns: # | Deliverable | Description | Acceptance Criteria
- Team/Roles → table with columns: Role | Name/Title | Responsibilities
- Payment Schedule → table with columns: Phase | Amount | Due Date | Trigger
- Any list of 3+ comparable items → convert to a table
- Tables ALWAYS have a heading above them and at least one sentence introducing them
- Table rows must be complete — never leave cells empty (use "N/A" or "TBD" if needed)

**Markdown Structure**
- Use H1 (#) for major section titles: "# 1. Executive Summary"
- Use H2 (##) for subsections: "## 1.1 Background"
- Use H3 (###) for sub-subsections: "### 1.1.1 Technical Context"
- Use bullet points for lists of 2-4 items
- Use numbered lists for sequential steps or ranked items
- Use **bold** for key terms, deliverable names, and deadlines
- Add horizontal rules (---) between major H1 sections for visual separation
- NEVER use horizontal rules between H2/H3 subsections — only between top-level sections

**Legal & Boilerplate Sections**
- Write standard professional SOW legal clauses (confidentiality, IP ownership, termination, dispute resolution)
- Use client name and project details from the proposals — NEVER use generic placeholders
- Keep legal language formal and precise

**Missing Data**
- If a detail is NOT in the proposals, mark it as [REVIEW NEEDED: specific thing missing]
- NEVER fabricate data to fill gaps — use [REVIEW NEEDED] instead
- Common fields that may need review: contact details, exact payment amounts, specific dates, signing authority names
- Do NOT make up milestone dates, budget figures, or team names that are not in the proposals

---

OUTPUT FORMAT:
- Pure markdown only — no preamble, no "Here is your SOW:", no closing remarks
- Start directly with the first heading of the SOW
- The document must be COMPREHENSIVE — a real SOW for a real engagement, not a summary"""


CHAT_SYSTEM = """You are a senior SOW editor at Nester Labs. You are refining a Statement of Work with the client in real time.

CURRENT SOW:
~~~
{sow_markdown}
~~~

{reference_context}

The SOW uses a FIXED 8-section Nester Labs structure. Do not add, remove, or reorder sections.
The visual design (Fira Sans, red accents, branded tables, side-by-side signatures) is hardcoded — you only change content.

EDITING STANDARDS:

**When making content changes:**
- Respond conversationally first: acknowledge the request, explain exactly what you changed and why
- Then output the COMPLETE updated SOW between ~~~SOW and ~~~ markers
- Never output just a diff or partial section — always the full document
- Maintain or improve the length — never shorten sections unless explicitly asked
- Keep the humanized, professional writing style throughout

**When asked to improve quality:**
- Expand thin sections with more detail, specific language, and professional prose
- Convert plain lists to tables where appropriate
- Replace vague language with specific terms, numbers, and commitments
- Remove AI-sounding filler phrases

**When asked about design/visual changes:**
- Explain that the Nester Labs brand design is permanently applied to all SOWs
- You can only modify the text content — visual styling is automatic

**Section 8 (Signatures):**
- NEVER write signature lines, underscores, or name/date fields in the SOW
- Only keep: "This SOW is executed by the duly authorized representatives of both parties as of the last signature date below."
- The actual signature block is programmatically generated in the DOCX

**When no changes needed:**
- Answer the question directly without outputting a SOW block

RULES:
- Always output the COMPLETE SOW when making content changes
- Preserve the exact 8-section structure — never merge, split, or reorder
- Keep [REVIEW NEEDED] markers unless the user specifically fills them in
- Be specific about every change made"""


def _load_docs_with_text(session_id: str | None, doc_type: str) -> list[dict[str, Any]]:
    """Load documents including extracted_text."""
    from memory.sqlite_ops import _get_conn
    conn = _get_conn()
    if session_id is None:
        rows = conn.execute(
            "SELECT id, file_name, extracted_text FROM sow_documents WHERE doc_type = ? AND session_id IS NULL",
            (doc_type,),
        ).fetchall()
    else:
        rows = conn.execute(
            "SELECT id, file_name, extracted_text FROM sow_documents WHERE doc_type = ? AND session_id = ?",
            (doc_type, session_id),
        ).fetchall()
    return [dict(r) for r in rows]


def _collect_text(docs: list[dict[str, Any]], max_chars: int = 40_000) -> str:
    """Join extracted text from documents, respecting char budget."""
    parts: list[str] = []
    used = 0
    for d in docs:
        text = d.get("extracted_text", "")
        if used + len(text) > max_chars:
            text = text[: max_chars - used] + "\n...[truncated]"
        parts.append(f"--- {d.get('file_name', 'document')} ---\n{text}")
        used += len(text)
        if used >= max_chars:
            break
    return "\n\n".join(parts)


async def _stream_generate(session_id: str):
    """Stream initial SOW generation."""
    from config.models import get_model, build_chat_llm
    from memory.sqlite_ops import update_sow_markdown, save_sow_chat_message

    # Higher temperature for more natural, humanized writing; max tokens for full-length output
    llm = build_chat_llm(get_model("synthesis"), temperature=0.7, max_tokens=16000)

    proposals = _load_docs_with_text(session_id, "proposal")
    proposal_text = _collect_text(proposals, max_chars=60_000)

    if proposal_text:
        user_msg = (
            "=== CLIENT PROPOSALS ===\n"
            "Extract every detail: names, dates, numbers, deliverables, milestones, requirements, "
            "client name, contact info, billing terms, assumptions.\n"
            "Use this as the sole source of new content. Be exhaustive — do not compress or skip anything.\n\n"
            + proposal_text
            + "\n\nNow generate the complete, full-length SOW following the fixed section structure."
        )
    else:
        user_msg = (
            "No client proposals were provided. Generate a comprehensive, professional SOW template "
            "for Nester Labs with all 8 sections fully written out. Use realistic placeholder content "
            "and mark client-specific fields with [REVIEW NEEDED: reason]."
        )

    messages = [
        {"role": "system", "content": GENERATE_SYSTEM},
        {"role": "user", "content": user_msg},
    ]

    full_response = ""
    try:
        async for chunk in llm.astream(messages):
            token = chunk.content
            if token:
                full_response += token
                yield f"data: {json.dumps({'token': token})}\n\n"

        update_sow_markdown(session_id, full_response)
        save_sow_chat_message(session_id, "assistant", full_response)

        yield f"data: {json.dumps({'sow_updated': True})}\n\n"
        yield "data: [DONE]\n\n"

    except Exception as exc:
        logger.error("[SOW] Generate stream error: %s", exc)
        yield f"data: {json.dumps({'error': str(exc)})}\n\n"
        yield "data: [DONE]\n\n"


@router.post("/sessions/{session_id}/generate")
async def generate_sow(session_id: str):
    from memory.sqlite_ops import get_sow_session, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    if not get_sow_session(session_id):
        raise HTTPException(404, "Session not found")

    return StreamingResponse(
        _stream_generate(session_id),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── SOW Chat (SSE) ──────────────────────────────────────────────────────────


async def _stream_chat(session_id: str, user_message: str):
    """Stream chat refinement and update SOW if changed."""
    from config.models import get_model, build_chat_llm
    from memory.sqlite_ops import (
        get_sow_session, update_sow_markdown,
        save_sow_chat_message, list_sow_chat_messages,
    )

    llm = build_chat_llm(get_model("synthesis"), temperature=0.3)

    session = get_sow_session(session_id)
    sow_md = session["sow_markdown"] if session else ""

    proposals = _load_docs_with_text(session_id, "proposal")

    ref_context = ""
    if proposals:
        ref_context += f"PROPOSAL CONTEXT:\n{_collect_text(proposals, 15_000)}\n\n"

    system = CHAT_SYSTEM.format(sow_markdown=sow_md, reference_context=ref_context)

    history = list_sow_chat_messages(session_id)
    messages: list[dict[str, str]] = [{"role": "system", "content": system}]
    for msg in history[-20:]:
        content = msg["content"]
        if msg["role"] == "assistant":
            content = re.sub(r"~~~SOW[\s\S]*?~~~", "[SOW content — see current version above]", content)
        messages.append({"role": msg["role"], "content": content})
    messages.append({"role": "user", "content": user_message})

    save_sow_chat_message(session_id, "user", user_message)

    full_response = ""
    try:
        async for chunk in llm.astream(messages):
            token = chunk.content
            if token:
                full_response += token
                yield f"data: {json.dumps({'token': token})}\n\n"

        save_sow_chat_message(session_id, "assistant", full_response)

        sow_match = re.search(r"~~~SOW\s*\n([\s\S]*?)\n~~~", full_response)
        if sow_match:
            new_sow = sow_match.group(1).strip()
            if new_sow:
                update_sow_markdown(session_id, new_sow)
                yield f"data: {json.dumps({'sow_updated': True})}\n\n"

        yield "data: [DONE]\n\n"

    except Exception as exc:
        logger.error("[SOW] Chat stream error: %s", exc)
        yield f"data: {json.dumps({'error': str(exc)})}\n\n"
        yield "data: [DONE]\n\n"


@router.post("/sessions/{session_id}/chat")
async def sow_chat(session_id: str, req: SOWChatReq):
    from memory.sqlite_ops import get_sow_session, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    if not get_sow_session(session_id):
        raise HTTPException(404, "Session not found")

    return StreamingResponse(
        _stream_chat(session_id, req.message),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── Direct Edit ──────────────────────────────────────────────────────────────


@router.put("/sessions/{session_id}/sow")
async def update_sow(session_id: str, req: UpdateSOWReq):
    from memory.sqlite_ops import update_sow_markdown, get_sow_session, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    if not get_sow_session(session_id):
        raise HTTPException(404, "Session not found")
    update_sow_markdown(session_id, req.markdown)
    return {"ok": True}


# ── DOCX Builder ─────────────────────────────────────────────────────────────


# ── Nester Brand Design System ──────────────────────────────────────────────
# Adapted from the ReportLab PDF design spec for python-docx DOCX output.

BRAND = {
    "colors": {
        "primary_red": "C0180C",
        "black": "181818",
        "dark_grey": "505050",
        "mid_grey": "707070",
        "light_grey": "E0E0E0",
        "table_header_fill": "E8E8E8",
    },
    "fonts": {
        # python-docx uses the Windows font name; Fira Sans must be installed.
        # Fallback chain: Fira Sans → Calibri → Arial
        "body": "Fira Sans Light",
        "body_fallback": "Calibri",
        "semibold": "Fira Sans SemiBold",
        "medium": "Fira Sans Medium",
        "bold": "Fira Sans Bold",       # cover only
        "light_italic": "Fira Sans Light Italic",
    },
    "sizes": {
        "h1": 22,      # section titles — SemiBold, NOT Bold
        "h2": 14,      # subsection titles — SemiBold
        "h3": 11,      # sub-subsection — Medium
        "h4": 10,      # minor heading — Medium
        "body": 9.5,   # body text — Light
        "table_header": 8,
        "table_body": 8,
        "footer": 6.5,
    },
    "spacing": {
        "before_h1": 18,   # pt before major heading
        "after_h1": 0,    # zero gap — red rule sits directly under title
        "before_h2": 14,
        "after_h2": 4,
        "before_h3": 10,
        "after_h3": 3,
        "body_after": 4,
        "bullet_after": 2,
        "table_row_height": 22,
    },
}


def _find_style(doc: Any, names: list[str], fallback: str = "Normal") -> str:
    """Find the first matching style name that exists in the document."""
    available = {s.name for s in doc.styles}
    for name in names:
        if name in available:
            return name
    return fallback


def _safe_add_paragraph(doc: Any, text: str, style: str) -> Any:
    """Add a paragraph with the given style, falling back to Normal on error."""
    try:
        return doc.add_paragraph(text, style=style)
    except Exception:
        return doc.add_paragraph(text)


def _markdown_to_docx(markdown_text: str) -> bytes:
    """Convert markdown to a branded Nester DOCX with hardcoded design system."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    b = BRAND
    colors = b["colors"]
    fonts = b["fonts"]
    sizes = b["sizes"]
    spacing = b["spacing"]

    # Brand values — always used, no overrides
    body_font = fonts["body"]
    body_size = sizes["body"]
    body_color = colors["dark_grey"]

    h1_font = fonts["semibold"]
    h1_size = sizes["h1"]
    h1_color = colors["black"]

    h2_font = fonts["semibold"]
    h2_size = sizes["h2"]
    h2_color = colors["black"]

    h3_font = fonts["medium"]
    h3_size = sizes["h3"]
    h3_color = colors["dark_grey"]

    doc = Document()

    # Page margins (match brand: ML=58pt, MR=50pt, MT=52pt, MB=66pt)
    for section in doc.sections:
        section.left_margin = Emu(int(58 * 12700))
        section.right_margin = Emu(int(50 * 12700))
        section.top_margin = Emu(int(52 * 12700))
        section.bottom_margin = Emu(int(66 * 12700))

    # --- Apply styles ---
    available_styles = {s.name for s in doc.styles}

    # Normal / body
    try:
        ns = doc.styles["Normal"]
        ns.font.name = body_font
        ns.font.size = Pt(body_size)
        ns.font.color.rgb = RGBColor.from_string(body_color)
        pf = ns.paragraph_format
        pf.space_after = Pt(spacing["body_after"])
        pf.line_spacing = 1.4
    except Exception:
        pass

    # Heading styles
    heading_specs = [
        ("Heading 1", h1_font, h1_size, h1_color, False, spacing["before_h1"], spacing["after_h1"]),
        ("Heading 2", h2_font, h2_size, h2_color, False, spacing["before_h2"], spacing["after_h2"]),
        ("Heading 3", h3_font, h3_size, h3_color, False, spacing["before_h3"], spacing["after_h3"]),
        ("Heading 4", h3_font, sizes["h4"], h3_color, False, 8, 3),
    ]
    for style_name, font, size, color, bold, sp_before, sp_after in heading_specs:
        if style_name not in available_styles:
            continue
        try:
            hs = doc.styles[style_name]
            hs.font.name = font
            hs.font.size = Pt(size)
            hs.font.color.rgb = RGBColor.from_string(color)
            hs.font.bold = bold  # SemiBold via font name, not bold flag
            hs.paragraph_format.space_before = Pt(sp_before)
            hs.paragraph_format.space_after = Pt(sp_after)
        except Exception:
            pass

    # Resolve style names (with fallbacks)
    h1_style = _find_style(doc, ["Heading 1", "heading 1", "Title", "Normal"])
    h2_style = _find_style(doc, ["Heading 2", "heading 2", "Subtitle", "Normal"])
    h3_style = _find_style(doc, ["Heading 3", "heading 3", "Normal"])
    h4_style = _find_style(doc, ["Heading 4", "heading 4", "Normal"])
    bullet_style = _find_style(doc, ["List Bullet", "List Bullet 1", "Bullet List", "Normal"])
    number_style = _find_style(doc, ["List Number", "List Number 1", "Numbered List", "Normal"])
    normal_style = _find_style(doc, ["Normal", "Body Text", "Body", "Default Paragraph Font"])

    # --- Helper: add red accent rule (H1 only) ---
    # Uses a real paragraph bottom-border instead of text characters
    def _add_red_rule() -> None:
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        # Add a bottom border to the paragraph itself (real line, not text)
        p_pr = p._p.get_or_add_pPr()
        p_bdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{colors["primary_red"]}"/>'
            f'</w:pBdr>'
        )
        p_pr.append(p_bdr)
        # Limit rule width via indentation — indent right by 70% of page width
        # Page content width ~487pt, so indent right ~340pt to get ~30% width line
        ind = parse_xml(
            f'<w:ind {nsdecls("w")} w:right="6800"/>'
        )
        p_pr.append(ind)

    # --- Parse markdown lines ---
    lines = markdown_text.split("\n")
    i = 0
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        # Table detection
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if not all(set(c.strip()) <= {"-", ":", " "} for c in cells):
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            _add_branded_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Headings
        if line.startswith("#### "):
            p = _safe_add_paragraph(doc, line[5:].strip(), h4_style)
        elif line.startswith("### "):
            p = _safe_add_paragraph(doc, line[4:].strip(), h3_style)
        elif line.startswith("## "):
            p = _safe_add_paragraph(doc, line[3:].strip(), h2_style)
        elif line.startswith("# "):
            p = _safe_add_paragraph(doc, line[2:].strip(), h1_style)
            _add_red_rule()
        # Bullets
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            p = _safe_add_paragraph(doc, "", bullet_style)
            _add_branded_runs(p, line.strip()[2:], body_font, body_size, body_color)
        # Numbered
        elif re.match(r"^\s*\d+\.\s", line):
            p = _safe_add_paragraph(doc, "", number_style)
            _add_branded_runs(p, re.sub(r"^\s*\d+\.\s", "", line), body_font, body_size, body_color)
        # Horizontal rule — skip visible line, just add spacing
        elif line.strip() in ("---", "***", "___"):
            pass
        # Blank
        elif not line.strip():
            pass
        # Normal body text
        else:
            p = _safe_add_paragraph(doc, "", normal_style)
            _add_branded_runs(p, line, body_font, body_size, body_color)

        i += 1

    if in_table and table_rows:
        _add_branded_table(doc, table_rows)

    # ── Side-by-side signature block (Nester brand) ──
    _add_signature_block(doc, fonts, colors, Pt)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_branded_runs(
    paragraph: Any,
    text: str,
    body_font: str,
    body_size: float,
    body_color: str,
) -> None:
    """Add runs with brand-aware bold (SemiBold) and italic (Light Italic) formatting."""
    from docx.shared import Pt, RGBColor

    b = BRAND
    semibold = b["fonts"]["semibold"]
    light_italic = b["fonts"]["light_italic"]

    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Bold → SemiBold per brand spec (never true Bold for body)
            run = paragraph.add_run(part[2:-2])
            run.font.name = semibold
            run.font.size = Pt(body_size)
            run.font.color.rgb = RGBColor.from_string(b["colors"]["black"])
            run.bold = False  # weight comes from font name
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = light_italic
            run.font.size = Pt(body_size)
            run.font.color.rgb = RGBColor.from_string(body_color)
            run.italic = True
        elif part:
            run = paragraph.add_run(part)
            run.font.name = body_font
            run.font.size = Pt(body_size)
            run.font.color.rgb = RGBColor.from_string(body_color)


def _add_branded_table(doc: Any, rows: list[list[str]]) -> None:
    """Add a Nester-branded table: grey header fill, proper fonts, red accent on totals."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    if not rows:
        return

    b = BRAND
    colors = b["colors"]
    fonts = b["fonts"]
    sizes = b["sizes"]

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)

    # Try to set Table Grid style
    try:
        table.style = "Table Grid"
    except (KeyError, Exception):
        pass

    for ri, row in enumerate(rows):
        is_header = ri == 0
        is_total = ri == len(rows) - 1 and any(
            "total" in c.lower() for c in row if c
        )

        for ci, cell_text in enumerate(row):
            if ci >= col_count:
                continue
            cell = table.rows[ri].cells[ci]
            cell.text = ""

            # Header row: grey fill + vertical center
            if is_header:
                tc_pr = cell._tc.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["table_header_fill"]}"/>')
                tc_pr.append(shading)
                v_align = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
                tc_pr.append(v_align)

            p = cell.paragraphs[0]
            run = p.add_run(cell_text)

            if is_header:
                run.font.name = fonts["medium"]
                run.font.size = Pt(sizes["table_header"])
                run.font.color.rgb = RGBColor.from_string(colors["dark_grey"])
                run.bold = False
            elif is_total:
                run.font.name = fonts["semibold"]
                run.font.size = Pt(sizes["table_body"] + 0.5)
                # Last column of total row in red
                if ci == col_count - 1:
                    run.font.color.rgb = RGBColor.from_string(colors["primary_red"])
                else:
                    run.font.color.rgb = RGBColor.from_string(colors["black"])
            else:
                run.font.name = fonts["body"]
                run.font.size = Pt(sizes["table_body"])
                run.font.color.rgb = RGBColor.from_string(colors["dark_grey"])


def _add_signature_block(
    doc: Any,
    fonts: dict[str, str],
    colors: dict[str, str],
    Pt: Any,
) -> None:
    """Add a side-by-side signature block: Nester Labs on the left, Client on the right."""
    from docx.shared import RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Spacer before signatures
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(24)
    spacer.paragraph_format.space_after = Pt(6)

    # Use a 2-column table with invisible borders for side-by-side layout
    table = doc.add_table(rows=5, cols=2)
    table.autofit = True

    # Remove borders
    try:
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        tbl_pr = table._tbl.tblPr if table._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>'
        )
        tbl_pr.append(borders)
    except Exception:
        pass

    # Row data: (left_text, right_text, font_name, font_size, color)
    sig_rows = [
        ("FOR NESTERLABS INDIA PRIVATE LIMITED", "ACCEPTED BY CLIENT", fonts["semibold"], 6.5, colors["dark_grey"]),
        ("─" * 35, "─" * 35, fonts["body"], 7, "C8C8C8"),
        ("Kunal Srivastava", "[Client Name]", fonts["body"], 7.5, colors["dark_grey"]),
        ("Director", "[Title]", fonts["body"], 7, colors["mid_grey"]),
        ("Date: ___________________________", "Date: ___________________________", fonts["body"], 7, colors["dark_grey"]),
    ]

    for ri, (left, right, font, size, color) in enumerate(sig_rows):
        for ci, text in enumerate([left, right]):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = font
            run.font.size = Pt(size)
            try:
                run.font.color.rgb = RGBColor.from_string(color)
            except Exception:
                pass


# ── DOCX → HTML Preview ─────────────────────────────────────────────────────


def _docx_to_html(docx_bytes: bytes) -> str:
    """Convert DOCX to HTML with hardcoded Nester brand CSS for faithful preview."""
    import mammoth

    result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    body_html = result.value

    b = BRAND
    c = b["colors"]
    s = b["sizes"]

    body_font = "Fira Sans"
    body_size = s["body"]
    body_color = f"#{c['dark_grey']}"
    h1_color = f"#{c['black']}"
    h2_color = f"#{c['black']}"
    h3_color = f"#{c['dark_grey']}"
    h1_size = s["h1"]
    h2_size = s["h2"]
    h3_size = s["h3"]

    # Swiss baseline grid unit (all spacing is a multiple of this)
    grid = 6  # pt

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<link href="https://fonts.googleapis.com/css2?family=Fira+Sans:ital,wght@0,200;0,300;0,400;0,500;0,600;0,700;1,300&display=swap" rel="stylesheet">
<style>
  /* ══════════════════════════════════════════════════════════════
     SWISS / INTERNATIONAL TYPOGRAPHIC STYLE
     Baseline grid: {grid}pt — all vertical spacing snaps to this unit
     ══════════════════════════════════════════════════════════════ */

  @page {{ size: A4; margin: 0; }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}

  body {{
    font-family: '{body_font}', 'Fira Sans', Calibri, 'Segoe UI', sans-serif;
    font-weight: 300;
    font-size: {body_size}pt;
    line-height: {grid * 3}pt;            /* 18pt baseline grid */
    color: {body_color};
    background: #f0f0f0;
    padding: 24px;
  }}

  .page {{
    max-width: 8.27in;
    min-height: 11.69in;
    margin: 0 auto;
    padding: 52pt 50pt 66pt 58pt;        /* brand margins: MT MR MB ML */
    background: #ffffff;
    box-shadow: 0 2px 20px rgba(0,0,0,0.12);
    border-radius: 2px;
    position: relative;
  }}

  /* ── SWISS GRID: Headings ──────────────────────────────────── */
  /* Every heading snaps to the same vertical rhythm: 3 grid units
     above, red rule below, 1 grid unit gap to body content.
     This creates the uniform heading→para pattern across all sections. */

  h1 {{
    font-family: 'Fira Sans', sans-serif;
    font-weight: 600;
    font-size: {h1_size}pt;
    line-height: {grid * 5}pt;
    color: {h1_color};
    margin: {grid * 4}pt 0 0;
    padding: 0 0 2pt;
    letter-spacing: -0.3px;
    border-bottom: 2pt solid #{c['primary_red']};
    border-bottom-width: 2pt;
    margin-bottom: {grid}pt;
  }}
  /* Shorten the red rule to ~20% width via pseudo-element */
  h1 {{
    border-bottom: none;
  }}
  h1::after {{
    content: '';
    display: block;
    width: 20%;
    height: 3pt;
    background: #{c['primary_red']};
    margin-top: 1pt;
    margin-bottom: {grid}pt;
  }}

  h2 {{
    font-family: 'Fira Sans', sans-serif;
    font-weight: 600;
    font-size: {h2_size}pt;
    line-height: {grid * 4}pt;
    color: {h2_color};
    margin: {grid * 3}pt 0 {grid}pt;
  }}

  h3 {{
    font-family: 'Fira Sans', sans-serif;
    font-weight: 500;
    font-size: {h3_size}pt;
    line-height: {grid * 3}pt;            /* 18pt */
    color: {h3_color};
    margin: {grid * 2}pt 0 {grid - 2}pt; /* 12pt above, 4pt below */
  }}

  h4 {{
    font-family: 'Fira Sans', sans-serif;
    font-weight: 500;
    font-size: {body_size}pt;
    line-height: {grid * 3}pt;
    color: {h3_color};
    margin: {grid * 2}pt 0 {grid - 3}pt;
  }}

  /* ── SWISS GRID: Body paragraphs ───────────────────────────── */
  /* Uniform spacing: every paragraph gets exactly 1 grid unit (6pt)
     below it. This locks body text to the baseline grid. */

  p {{
    margin: 0 0 {grid}pt;
    line-height: {grid * 3}pt;            /* 18pt — matches body baseline */
  }}

  /* First paragraph after a heading — no extra top margin (heading's
     margin-bottom already provides the grid gap) */
  h1 + p, h2 + p, h3 + p, h4 + p {{
    margin-top: 0;
  }}

  /* ── SWISS GRID: Lists ─────────────────────────────────────── */
  ul, ol {{
    margin: {grid}pt 0 {grid}pt;
    padding-left: 22pt;
    line-height: {grid * 3}pt;
  }}
  li {{
    margin: 0 0 {grid - 2}pt;            /* 4pt between items */
  }}

  /* ── SWISS GRID: Tables ────────────────────────────────────── */
  table {{
    width: 100%;
    border-collapse: collapse;
    margin: {grid * 2}pt 0;              /* 12pt above and below */
    font-size: {s['table_body']}pt;
    line-height: {grid * 3}pt;
  }}
  th, td {{
    border: 0.5pt solid #{c['light_grey']};
    padding: {grid - 1}pt 8pt;
    text-align: left;
    vertical-align: top;
  }}
  /* Header row — vertically centered */
  thead th, thead td, tr:first-child th, tr:first-child td {{
    vertical-align: middle;
  }}
  /* Header row */
  thead tr, tr:first-child {{
    background: #{c['table_header_fill']};
  }}
  th, tr:first-child td {{
    font-weight: 500;
    font-size: {s['table_header']}pt;
    color: #{c['dark_grey']};
    letter-spacing: 0.2px;
    text-transform: none;
  }}
  /* Total / last row accent */
  tr:last-child td {{
    font-weight: 600;
    border-top: 1pt solid #{c['dark_grey']};
  }}

  /* ── Inline formatting ─────────────────────────────────────── */
  strong {{
    font-weight: 600;
    color: #{c['black']};
  }}
  em {{
    font-weight: 300;
    font-style: italic;
    color: #{c['mid_grey']};
  }}

  /* ── Horizontal rules (major section dividers only) ─────────── */
  hr {{
    border: none;
    border-top: 0.5pt solid #{c['light_grey']};
    margin: {grid * 3}pt 0;              /* 18pt — 3 grid units */
  }}

  /* ── Spine accent (brand element) ──────────────────────────── */
  .page::after {{
    content: 'REIMAGINING  INTELLIGENCE';
    position: absolute;
    right: -2pt;
    top: 75%;
    transform: rotate(-90deg) translateX(50%);
    transform-origin: right center;
    font-family: 'Fira Sans', sans-serif;
    font-weight: 500;
    font-size: 5.5pt;
    color: #{c['primary_red']};
    letter-spacing: 2px;
    opacity: 0.7;
  }}
</style>
</head>
<body>
<div class="page">
{body_html}
</div>
</body>
</html>"""


@router.get("/sessions/{session_id}/preview")
async def preview_sow(session_id: str):
    """Generate DOCX from current SOW markdown, convert to HTML, return for iframe."""
    from memory.sqlite_ops import get_sow_session, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    session = get_sow_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if not session.get("sow_markdown", "").strip():
        raise HTTPException(400, "No SOW content")

    try:
        docx_bytes = _markdown_to_docx(session["sow_markdown"])
        html = _docx_to_html(docx_bytes)
        return HTMLResponse(content=html)
    except Exception as exc:
        import traceback
        logger.error("[SOW] Preview error for %s: %s\n%s", session_id, exc, traceback.format_exc())
        raise HTTPException(500, f"Preview failed: {exc}")


# ── DOCX Download ────────────────────────────────────────────────────────────


@router.get("/sessions/{session_id}/download")
async def download_sow(session_id: str):
    from memory.sqlite_ops import get_sow_session, is_sqlite_ready
    if not is_sqlite_ready():
        raise HTTPException(503, "Database not ready")
    session = get_sow_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if not session.get("sow_markdown", "").strip():
        raise HTTPException(400, "No SOW content to download")

    try:
        docx_bytes = _markdown_to_docx(session["sow_markdown"])
    except Exception as exc:
        import traceback
        logger.error("[SOW] Download error for %s: %s\n%s", session_id, exc, traceback.format_exc())
        raise HTTPException(500, f"DOCX generation failed: {exc}")

    title = session.get("title", "SOW").replace('"', "")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="SOW_{title}.docx"'},
    )
